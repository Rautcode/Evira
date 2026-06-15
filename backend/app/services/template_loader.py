import os
from typing import Dict, Any, Optional
from docx import Document
from jinja2 import Environment, FileSystemLoader

TEMPLATES_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../templates'))

class TemplateLoader:
    def __init__(self, templates_dir: Optional[str] = None):
        self.templates_dir = templates_dir or TEMPLATES_DIR
        self.env = Environment(loader=FileSystemLoader(self.templates_dir))

    def load_and_render_html(self, template_name: str, context: Dict[str, Any]) -> str:
        template = self.env.get_template(template_name)
        return template.render(**context)

    def load_and_render_docx(self, template_name: str, context: Dict[str, Any], output_path: Optional[str] = None) -> str:
        path = os.path.join(self.templates_dir, template_name)
        doc = Document(path)
        for p in doc.paragraphs:
            for key, value in context.items():
                if f'{{{{{key}}}}}' in p.text:
                    p.text = p.text.replace(f'{{{{{key}}}}}', str(value))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context.items():
                        if f'{{{{{key}}}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{{key}}}}}', str(value))
        if not output_path:
            output_path = os.path.join(self.templates_dir, f'rendered_{template_name}')
        doc.save(output_path)
        return output_path
